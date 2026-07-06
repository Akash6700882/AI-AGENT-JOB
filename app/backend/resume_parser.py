"""
Resume Parser Module - Extracts structured data from PDF and DOCX resumes
"""
import re
import json
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict
import logging

logger = logging.getLogger(__name__)

TECH_SKILLS_DB = {
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "golang",
    "rust", "ruby", "php", "swift", "kotlin", "scala", "perl", "r", "matlab",
    "dart", "lua", "haskell", "clojure", "elixir", "erlang", "julia", "groovy",
    "objective-c", "sql", "bash", "powershell",
    "react", "vue", "vue.js", "angular", "svelte", "next.js", "nuxt", "gatsby",
    "html", "html5", "css", "css3", "sass", "less", "tailwind", "bootstrap",
    "webpack", "vite", "rollup", "parcel", "jquery",
    "node.js", "nodejs", "express", "django", "flask", "fastapi", "spring",
    "spring boot", "laravel", "rails", "asp.net", "graphql", "rest", "restful",
    "api", "microservices", "websocket", "grpc",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "cassandra",
    "dynamodb", "sqlite", "mariadb", "neo4j", "couchdb", "firebase",
    "amazon rds", "planetscale", "prisma", "sqlalchemy", "mongoose",
    "aws", "amazon web services", "azure", "gcp", "google cloud", "docker",
    "kubernetes", "k8s", "terraform", "ansible", "jenkins", "gitlab ci",
    "github actions", "circleci", "travisci", "prometheus", "grafana",
    "nginx", "apache", "linux", "ubuntu", "debian", "centos", "unix",
    "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
    "scipy", "matplotlib", "seaborn", "jupyter", "opencv", "nltk", "spacy",
    "huggingface", "transformers", "llm", "machine learning", "deep learning",
    "data science", "nlp", "computer vision", "reinforcement learning",
    "xgboost", "lightgbm", "catboost",
    "react native", "flutter", "ionic", "cordova", "android", "ios",
    "xamarin", "phonegap", "expo",
    "git", "github", "gitlab", "bitbucket", "jira", "confluence", "trello",
    "figma", "sketch", "adobe xd", "postman", "insomnia", "swagger",
    "openapi", "kafka", "rabbitmq", "celery", "memcached",
    "electron", "tauri", "pwa", "webassembly", "wasm", "three.js", "d3.js",
    "agile", "scrum", "kanban", "tdd", "ci/cd", "oauth", "jwt", "sso",
    "reactjs", "vuejs", "angularjs", "nodejs", "python3",
    "postgres", "mongo", "k8", "gke", "eks", "ecs",
    "rest api", "restful api", "web api",
}

_LIGATURE_MAP = {
    "\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
    "\ufb03": "ffi", "\ufb04": "ffl", "\ufb05": "st", "\ufb06": "st",
    "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
    "\u2013": "-", "\u2014": "-",
}

JOB_TITLES_DB = [
    "software engineer", "frontend developer", "backend developer",
    "full stack developer", "web developer", "mobile developer",
    "devops engineer", "data scientist", "machine learning engineer",
    "ai engineer", "cloud engineer", "site reliability engineer",
    "security engineer", "qa engineer", "test engineer",
    "product manager", "project manager", "scrum master",
    "ux designer", "ui designer", "product designer",
    "data engineer", "data analyst", "business analyst",
    "system administrator", "network engineer", "database administrator",
    "technical lead", "engineering manager", "cto", "vp of engineering",
    "solutions architect", "enterprise architect", "technical architect",
]


@dataclass
class ResumeData:
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    github: str = ""
    portfolio: str = ""
    summary: str = ""
    skills: List[str] = None
    experience: List[Dict] = None
    education: List[Dict] = None
    certifications: List[str] = None
    raw_text: str = ""

    def __post_init__(self):
        if self.skills is None:
            self.skills = []
        if self.experience is None:
            self.experience = []
        if self.education is None:
            self.education = []
        if self.certifications is None:
            self.certifications = []

    def to_dict(self) -> Dict:
        return asdict(self)


class ResumeParser:
    def __init__(self):
        self.extracted_text = ""

    def parse_file(self, file_path: str) -> ResumeData:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        if path.suffix.lower() == ".pdf":
            self.extracted_text = self._parse_pdf(file_path)
        elif path.suffix.lower() in [".docx", ".doc"]:
            self.extracted_text = self._parse_docx(file_path)
        elif path.suffix.lower() == ".txt":
            self.extracted_text = path.read_text(encoding="utf-8")
        else:
            raise ValueError(f"Unsupported file format: {path.suffix}")

        if not self.extracted_text or not self.extracted_text.strip():
            logger.warning("No text extracted from %s", file_path)
            return ResumeData()

        logger.info("Extracted %d characters from resume", len(self.extracted_text))

        resume = ResumeData(raw_text=self.extracted_text)
        resume.name       = self._extract_name()
        resume.email      = self._extract_email()
        resume.phone      = self._extract_phone()
        resume.location   = self._extract_location()
        resume.linkedin   = self._extract_linkedin()
        resume.github     = self._extract_github()
        resume.skills     = self._extract_skills()
        resume.experience = self._extract_experience()
        resume.education  = self._extract_education()
        resume.summary    = self._extract_summary()

        logger.info("Skills found (%d): %s", len(resume.skills), resume.skills)
        return resume

    def _parse_pdf(self, file_path: str) -> str:
        # Try pdfplumber first — much better extraction
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            text = "\n".join(parts)
            if text.strip():
                return text
            logger.warning("pdfplumber returned empty, trying PyPDF2")
        except ImportError:
            logger.info("pdfplumber not installed, trying PyPDF2")
        except Exception as e:
            logger.warning("pdfplumber failed: %s", e)

        try:
            import PyPDF2
            parts = []
            with open(file_path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            return "\n".join(parts)
        except ImportError:
            raise ImportError("Install a PDF library:  pip install pdfplumber")

    def _parse_docx(self, file_path: str) -> str:
        try:
            import docx
            doc = docx.Document(file_path)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except ImportError:
            raise ImportError("Run:  pip install python-docx")

    def _normalize_text(self, text: str) -> str:
        for char, rep in _LIGATURE_MAP.items():
            text = text.replace(char, rep)
        text = re.sub(r"[ \t]+", " ", text)
        return text

    def _build_skill_pattern(self, skill: str) -> str:
        escaped = re.escape(skill)
        escaped = escaped.replace(r"\ ", r"\s+")
        prefix = r"\b"  if re.match(r"\w", skill[0])  else r"(?<!\w)"
        suffix = r"\b"  if re.match(r"\w", skill[-1]) else r"(?!\w)"
        return prefix + escaped + suffix

    def _extract_skills(self) -> List[str]:
        if not self.extracted_text:
            return []
        text_lower = self._normalize_text(self.extracted_text).lower()
        found: set = set()
        for skill in TECH_SKILLS_DB:
            try:
                if re.search(self._build_skill_pattern(skill), text_lower):
                    found.add(skill)
            except re.error as e:
                logger.warning("Bad pattern for %r: %s", skill, e)
        return sorted(found)

    def _split_sections(self) -> Dict[str, str]:
        header_re = re.compile(
            r"^\s*(experience|work history|employment|professional experience|"
            r"education|academic background|skills|technical skills|"
            r"projects|certifications|summary|profile|objective)\s*$",
            re.I | re.M,
        )
        sections: Dict[str, str] = {}
        current_key = "header"
        current_lines: List[str] = []
        for line in self.extracted_text.split("\n"):
            m = header_re.match(line)
            if m:
                sections[current_key] = "\n".join(current_lines).strip()
                current_key = m.group(1).lower()
                current_lines = []
            else:
                current_lines.append(line)
        sections[current_key] = "\n".join(current_lines).strip()
        return sections

    def _extract_name(self) -> str:
        skip = re.compile(
            r"(resume|curriculum vitae|\bcv\b|summary|objective|profile|"
            r"experience|education|skills|contact|phone|email|address|"
            r"linkedin|github|http|www\.|\d{5})", re.I,
        )
        for line in self.extracted_text.strip().split("\n")[:8]:
            line = line.strip()
            if not line or len(line) > 60 or "@" in line:
                continue
            if skip.search(line):
                continue
            tokens = line.split()
            if 2 <= len(tokens) <= 4 and all(t[0].isupper() for t in tokens if t.isalpha()):
                return line
        return ""

    def _extract_email(self) -> str:
        m = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
                      self.extracted_text)
        return m.group(0) if m else ""

    def _extract_phone(self) -> str:
        m = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
                      self.extracted_text)
        return m.group(0) if m else ""

    def _extract_location(self) -> str:
        for pat in [r"([A-Za-z\s]+,\s*[A-Za-z]{2}\s*\d{5})",
                    r"([A-Za-z\s]+,\s*[A-Za-z\s]{2,})"]:
            m = re.search(pat, self.extracted_text)
            if m:
                return m.group(0).strip()
        return ""

    def _extract_linkedin(self) -> str:
        m = re.search(r"linkedin\.com/in/[A-Za-z0-9_%-]+", self.extracted_text, re.I)
        return f"https://www.{m.group(0)}" if m else ""

    def _extract_github(self) -> str:
        m = re.search(r"github\.com/[A-Za-z0-9_-]+", self.extracted_text, re.I)
        return f"https://www.{m.group(0)}" if m else ""

    def _extract_experience(self) -> List[Dict]:
        sections = self._split_sections()
        exp_text = ""
        for key in ("experience", "work history", "employment", "professional experience"):
            if key in sections and sections[key]:
                exp_text = sections[key]
                break
        if not exp_text:
            return []
        experience = []
        entries = re.split(r"\n(?=[A-Z][A-Za-z ]{3,60}\n)", exp_text)
        for entry in entries:
            lines = [l.strip() for l in entry.strip().split("\n") if l.strip()]
            if len(lines) >= 2:
                experience.append({
                    "title": lines[0],
                    "company": lines[1],
                    "description": " ".join(lines[2:])[:300] if len(lines) > 2 else "",
                })
        return experience

    def _extract_education(self) -> List[Dict]:
        sections = self._split_sections()
        edu_text = ""
        for key in ("education", "academic background"):
            if key in sections and sections[key]:
                edu_text = sections[key]
                break
        if not edu_text:
            return []
        education = []
        degree_re = re.compile(
            r"(B\.?S\.?|M\.?S\.?|Ph\.?D\.?|Bachelor[^\n,]*|Master[^\n,]*|"
            r"MBA|B\.?E\.?|B\.?Tech\.?|M\.?Tech\.?|Associate[^\n,]*)", re.I
        )
        year_re = re.compile(r"\b(19|20)\d{2}\b")
        for line in edu_text.split("\n"):
            dm = degree_re.search(line)
            if dm:
                ym = year_re.search(line)
                education.append({
                    "degree": dm.group(0).strip(),
                    "field": "", "institution": "",
                    "year": ym.group(0) if ym else "",
                })
        return education

    def _extract_summary(self) -> str:
        sections = self._split_sections()
        for key in ("summary", "profile", "objective"):
            text = sections.get(key, "")
            if text and len(text) > 20:
                return text[:500]
        return ""

    def get_skills_vector(self) -> Dict[str, float]:
        return {skill: 1.0 for skill in self._extract_skills()}